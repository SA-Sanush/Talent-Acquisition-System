import os
import re
import csv
import json
import math
import shutil
import sqlite3
import secrets
import traceback
import uuid
from ast import literal_eval
from collections import Counter
from datetime import datetime, timezone
from functools import wraps
from io import BytesIO
from pathlib import Path
import numpy as np
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, flash, session, abort, Response
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

try:
    from pypdf import PdfReader
except ImportError:  # Backward-compatible fallback for older local environments.
    from PyPDF2 import PdfReader

try:
    import spacy
except ImportError:  # Optional: app still works with the local rule fallback.
    spacy = None

try:
    from docx import Document
except ImportError:  # Optional until python-docx is installed.
    Document = None

OCR_IMPORT_ERRORS: dict[str, str] = {}

try:
    import pytesseract
except ImportError as exc:  # OCR is optional because it needs native Tesseract binaries.
    pytesseract = None
    OCR_IMPORT_ERRORS["pytesseract"] = str(exc)

try:
    from pdf2image import convert_from_path
except ImportError as exc:  # OCR is optional because it needs native Poppler binaries.
    convert_from_path = None
    OCR_IMPORT_ERRORS["pdf2image"] = str(exc)

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_FOLDER = Path(os.environ.get("UPLOAD_FOLDER", str(BASE_DIR / "uploads")))
CHART_FOLDER = BASE_DIR / "static" / "images"
JOB_DATA_PATH = BASE_DIR / "job_data.csv"
DB_PATH = Path(os.environ.get("DATABASE_PATH", str(BASE_DIR / "tas_reports.db")))
ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}
RATE_LIMITS: dict[tuple[str, str], list[float]] = {}
_DB_READY = False
LAST_OCR_ERROR = ""

COMMON_TESSERACT_PATHS = [
    Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
    Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"),
]
COMMON_POPPLER_PATHS = [
    Path(r"C:\Program Files\poppler\Library\bin"),
    Path(r"C:\Program Files\poppler\bin"),
]


def path_exists(path: Path) -> bool:
    try:
        return path.exists()
    except OSError:
        return False


def configured_tesseract_path() -> str:
    if pytesseract is None:
        return ""
    candidates = [
        os.environ.get("TESSERACT_CMD", ""),
        getattr(pytesseract.pytesseract, "tesseract_cmd", ""),
        *(str(path) for path in COMMON_TESSERACT_PATHS),
    ]
    for candidate in candidates:
        if candidate and path_exists(Path(candidate)):
            return candidate
    found = shutil.which("tesseract")
    return found or ""


def configure_tesseract_command() -> None:
    if pytesseract is None:
        return
    command = configured_tesseract_path()
    if command:
        pytesseract.pytesseract.tesseract_cmd = command


configure_tesseract_command()


def winget_poppler_paths() -> list[Path]:
    package_root = Path.home() / "AppData" / "Local" / "Microsoft" / "WinGet" / "Packages"
    if not path_exists(package_root):
        return []
    try:
        package_dirs = list(package_root.glob("oschwartz10612.Poppler_*"))
    except OSError:
        return []
    paths: list[Path] = []
    for package_dir in package_dirs:
        try:
            paths.extend(package_dir.glob("poppler-*/*/bin"))
        except OSError:
            continue
    return paths


def configured_poppler_path() -> str:
    candidates = [
        os.environ.get("POPPLER_PATH", ""),
        *(str(path) for path in COMMON_POPPLER_PATHS),
        *(str(path) for path in winget_poppler_paths()),
    ]
    for candidate in candidates:
        if not candidate:
            continue
        candidate_path = Path(candidate)
        if path_exists(candidate_path / "pdftoppm.exe") or path_exists(candidate_path / "pdftoppm"):
            return candidate
    found = shutil.which("pdftoppm")
    return str(Path(found).parent) if found else ""

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY") or secrets.token_hex(32)
app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["CHART_FOLDER"] = str(CHART_FOLDER)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024
app.config["CSRF_ENABLED"] = os.environ.get("DISABLE_CSRF") != "1"
app.config["ADMIN_USERNAME"] = os.environ.get("ADMIN_USERNAME", "admin")
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.environ.get("FLASK_ENV") == "production",
)

UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)
CHART_FOLDER.mkdir(parents=True, exist_ok=True)

CONTACT_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}|\+?\d[\d\s().-]{6,}\d")
EMAIL_PATTERN = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$")
YEAR_PATTERN = re.compile(r"\b(19|20)\d{2}\b")
MAX_TEXT_INPUT_LENGTH = 5000
MAX_CATALOG_FIELD_LENGTH = 800

COMMON_TYPO = ["teh", "recieve", "acheive", "managment", "objecive", "formaing", "resposible", "seperated"]

JOB_KEYWORDS = [
    "communication", "problem solving", "technical", "leadership", "python", "sql", "flask",
    "data analysis", "teamwork", "project", "documentation", "presentation", "analytics",
    "stakeholders", "roadmap", "testing", "bug", "portfolio", "linkedin"
]

ACTION_VERBS = [
    "achieved", "analyzed", "automated", "built", "created", "delivered", "designed",
    "developed", "improved", "increased", "launched", "led", "managed", "optimized",
    "reduced", "shipped", "streamlined"
]

WEAK_VERBS = ["helped", "worked on", "responsible for", "handled", "assisted", "participated"]

SKILL_SYNONYMS = {
    "python": ["python", "py"],
    "sql": ["sql", "mysql", "postgresql", "postgres", "sqlite", "database"],
    "pandas": ["pandas", "dataframe", "dataframes"],
    "excel": ["excel", "spreadsheets", "vlookup", "pivot table"],
    "data visualization": ["data visualization", "tableau", "power bi", "matplotlib", "seaborn", "dashboard"],
    "machine learning": ["machine learning", "ml", "machine-learning", "scikit-learn", "sklearn", "modeling", "predictive"],
    "nlp": ["nlp", "natural language processing", "spacy", "bert", "transformers"],
    "flask": ["flask", "web api", "rest api", "rest apis"],
    "javascript": ["javascript", "js", "typescript", "react", "node"],
    "api development": ["api", "apis", "rest api", "rest apis", "endpoint", "endpoints"],
    "git": ["git", "github", "version control"],
    "html": ["html", "html5"],
    "css": ["css", "css3", "responsive design", "responsive"],
    "communication": ["communication", "presentation", "stakeholder", "stakeholders", "collaboration"],
    "leadership": ["leadership", "led", "managed", "mentored", "ownership"],
    "testing": ["testing", "pytest", "unit test", "qa", "automation"],
}

ROLE_KEYWORD_HINTS = {
    "full stack developer": [
        "javascript", "react", "node", "python", "flask", "sql", "git", "testing",
        "api", "html", "css"
    ],
    "frontend developer": ["javascript", "react", "typescript", "html", "css", "testing", "git"],
    "front end developer": ["javascript", "react", "typescript", "html", "css", "testing", "git"],
    "backend developer": ["python", "flask", "sql", "api", "testing", "git"],
    "back end developer": ["python", "flask", "sql", "api", "testing", "git"],
    "software engineer": ["python", "javascript", "sql", "git", "testing", "code review", "api"],
    "web developer": ["javascript", "html", "css", "react", "api", "git", "testing"],
    "data analyst": ["python", "sql", "excel", "data visualization", "analytics", "reporting"],
}

LOW_SIGNAL_ROLE_WORDS = {
    "role", "job", "candidate", "skills", "skill", "developer", "engineer", "full",
    "stack", "front", "frontend", "backend", "back", "end", "web", "software",
    "senior", "junior", "lead", "manager", "specialist", "analyst"
}

SCORING_WEIGHTS = {
    "Contact Information": 0.08,
    "Professional Summary": 0.10,
    "Work Experience": 0.18,
    "Education": 0.10,
    "Skills": 0.20,
    "TAS Optimization": 0.14,
    "Consistency": 0.07,
    "Proofreading": 0.06,
    "File Format": 0.04,
    "Relevance": 0.03,
}

ROLE_WEIGHT_OVERRIDES = {
    "data analyst": {"Skills": 0.28, "Relevance": 0.08, "Work Experience": 0.16, "Education": 0.08},
    "software engineer": {"Skills": 0.26, "Work Experience": 0.22, "TAS Optimization": 0.12, "Relevance": 0.06},
    "developer": {"Skills": 0.28, "Work Experience": 0.22, "TAS Optimization": 0.12, "Relevance": 0.06},
    "product manager": {"Professional Summary": 0.14, "Work Experience": 0.22, "Skills": 0.16, "Relevance": 0.08},
    "technical writer": {"Professional Summary": 0.16, "Proofreading": 0.12, "TAS Optimization": 0.12, "Relevance": 0.08},
}

try:
    NLP = spacy.load(os.environ.get("SPACY_MODEL", "en_core_web_sm")) if spacy else None
except OSError:
    NLP = spacy.blank("en") if spacy else None

DEFAULT_JOBS = [
    {
        "title": "Data Analyst",
        "keywords": "python, sql, pandas, data visualization, problem solving, communication, reporting",
        "description": "Analyze datasets, create dashboards, and deliver insights to business stakeholders."
    },
    {
        "title": "Software Engineer",
        "keywords": "python, javascript, flask, sql, git, testing, agile, code review",
        "description": "Design and implement applications, maintain code quality, and collaborate on engineering teams."
    },
    {
        "title": "Product Manager",
        "keywords": "communication, roadmap, stakeholders, analytics, strategy, prioritization",
        "description": "Define product direction, coordinate teams, and shape customer-focused roadmaps."
    },
    {
        "title": "Technical Writer",
        "keywords": "documentation, editing, communication, clarity, process, technical writing",
        "description": "Produce clear technical guides, manuals, and user-facing content for engineering teams."
    },
    {
        "title": "Digital Marketing Specialist",
        "keywords": "seo, content, analytics, copywriting, branding, social media, campaign",
        "description": "Create campaigns, analyze performance, and develop customer-facing digital content."
    }
]


def init_db() -> None:
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(
            "CREATE TABLE IF NOT EXISTS reports ("
            "id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "user_id INTEGER,"
            "filename TEXT,"
            "uploaded_at TEXT,"
            "score INTEGER,"
            "matched_jobs TEXT,"
            "criteria TEXT,"
            "suggestions TEXT,"
            "analysis TEXT"
            ")"
        )
        conn.execute(
            "CREATE TABLE IF NOT EXISTS contacts ("
            "id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "user_id INTEGER,"
            "name TEXT,"
            "email TEXT,"
            "subject TEXT,"
            "message TEXT,"
            "created_at TEXT"
            ")"
        )
        conn.execute(
            "CREATE TABLE IF NOT EXISTS users ("
            "id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "username TEXT UNIQUE,"
            "password_hash TEXT,"
            "created_at TEXT"
            ")"
        )
        conn.commit()
        ensure_user_columns(conn)
        ensure_contact_columns(conn)
        ensure_report_columns(conn)
        ensure_reports_user_scope(conn)
        create_default_user(conn)
        assign_legacy_reports_to_default_user(conn)
    finally:
        conn.close()


def create_default_user(conn: sqlite3.Connection) -> None:
    cursor = conn.execute("SELECT 1 FROM users LIMIT 1")
    if cursor.fetchone() is None:
        password_hash = generate_password_hash("Password123!")
        conn.execute(
            "INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
            ("admin", password_hash, utc_timestamp()),
        )
        conn.commit()


def get_primary_history_user_row(conn: sqlite3.Connection) -> tuple[int, str] | None:
    row = conn.execute(
        "SELECT id, username FROM users WHERE username <> ? ORDER BY id ASC LIMIT 1",
        ("admin",),
    ).fetchone()
    if row:
        return row
    return conn.execute(
        "SELECT id, username FROM users WHERE username = ? LIMIT 1",
        ("admin",),
    ).fetchone()


def parse_report_payload(value: str | None, fallback):
    if not value:
        return fallback
    try:
        return json.loads(value)
    except (json.JSONDecodeError, TypeError):
        try:
            return literal_eval(value)
        except (ValueError, SyntaxError):
            return fallback


def summarize_matched_jobs(matched_jobs_value: str | None) -> str:
    parsed = parse_report_payload(matched_jobs_value, [])
    if isinstance(parsed, list):
        titles = []
        for item in parsed:
            if isinstance(item, dict):
                title = item.get("title")
                if title:
                    titles.append(title)
            elif isinstance(item, str) and item.strip():
                titles.append(item.strip())
        return ", ".join(titles) if titles else "None"
    if isinstance(parsed, str):
        return parsed
    return "None"


def get_accessible_user_ids(conn: sqlite3.Connection, current_user_id: int) -> list[int]:
    accessible_user_ids = [current_user_id]
    primary_row = get_primary_history_user_row(conn)
    admin_row = conn.execute(
        "SELECT id FROM users WHERE username = ? LIMIT 1",
        ("admin",),
    ).fetchone()
    if primary_row and current_user_id == primary_row[0] and admin_row and admin_row[0] != current_user_id:
        accessible_user_ids.append(admin_row[0])
    return accessible_user_ids


def load_job_data() -> pd.DataFrame:
    if not JOB_DATA_PATH.exists():
        with open(JOB_DATA_PATH, "w", newline="", encoding="utf-8") as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=["title", "keywords", "description"])
            writer.writeheader()
            for row in DEFAULT_JOBS:
                writer.writerow(row)
    return pd.read_csv(JOB_DATA_PATH)


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def file_signature_matches(path: Path, extension: str) -> bool:
    try:
        header = path.read_bytes()[:8]
    except OSError:
        return False
    if extension == "pdf":
        return header.startswith(b"%PDF-")
    if extension == "docx":
        return header.startswith(b"PK\x03\x04")
    if extension == "txt":
        return b"\x00" not in header
    return False


def compact_reason(reason: str, limit: int = 150) -> str:
    reason = re.sub(r"\s+", " ", reason).strip()
    return reason if len(reason) <= limit else f"{reason[: limit - 3].rstrip()}..."


def count_bullet_markers(text: str) -> int:
    return text.count("- ") + text.count("\u2022") + text.count("* ")


def utc_now() -> datetime:
    return datetime.now(timezone.utc)


def utc_timestamp() -> str:
    return utc_now().isoformat(timespec="seconds")


def client_key() -> str:
    return request.headers.get("X-Forwarded-For", request.remote_addr or "local").split(",")[0].strip()


def rate_limit(limit: int, seconds: int):
    def decorator(view_func):
        @wraps(view_func)
        def wrapped(*args, **kwargs):
            if request.method == "GET":
                return view_func(*args, **kwargs)
            now = utc_now().timestamp()
            key = (view_func.__name__, client_key())
            recent = [stamp for stamp in RATE_LIMITS.get(key, []) if now - stamp < seconds]
            if len(recent) >= limit:
                flash("Too many attempts. Please wait a moment and try again.", "error")
                return redirect(request.referrer or url_for("index"))
            recent.append(now)
            RATE_LIMITS[key] = recent
            return view_func(*args, **kwargs)
        return wrapped
    return decorator


def csrf_token() -> str:
    token = session.get("_csrf_token")
    if not token:
        token = secrets.token_urlsafe(32)
        session["_csrf_token"] = token
    return token


@app.context_processor
def inject_csrf_token():
    return {"csrf_token": csrf_token}


@app.before_request
def protect_from_csrf() -> None:
    if not app.config.get("CSRF_ENABLED", True) or request.method not in {"POST", "PUT", "PATCH", "DELETE"}:
        return
    session_token = session.get("_csrf_token")
    form_token = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token")
    if not session_token or not form_token or not secrets.compare_digest(session_token, form_token):
        abort(400)


def save_upload_temporarily(uploaded_file, original_filename: str, user_id: int | None) -> Path:
    extension = original_filename.rsplit(".", 1)[1].lower()
    user_folder = UPLOAD_FOLDER / (str(user_id) if user_id else "anonymous")
    user_folder.mkdir(parents=True, exist_ok=True)
    save_path = user_folder / f"{uuid.uuid4().hex}.{extension}"
    uploaded_file.save(save_path)
    return save_path


def extract_text_from_pdf(path: Path) -> str:
    text_parts = []
    try:
        reader = PdfReader(str(path))
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
    except Exception:
        return ""
    return "\n".join(text_parts)


def ocr_unavailable_message() -> str:
    if pytesseract is None or convert_from_path is None:
        missing_packages = []
        if pytesseract is None:
            missing_packages.append("pytesseract")
        if convert_from_path is None:
            missing_packages.append("pdf2image")
        package_list = " and ".join(missing_packages)
        return (
            f"OCR Python package{'s are' if len(missing_packages) > 1 else ' is'} missing: "
            f"{package_list}. Install the missing package{'s' if len(missing_packages) > 1 else ''} "
            "with pip, then install the Tesseract OCR and Poppler system tools."
        )
    if not tesseract_is_available():
        return (
            "Tesseract OCR is not installed or not on PATH. Install Tesseract OCR, "
            "or set TESSERACT_CMD to the full path of tesseract.exe."
        )
    if not poppler_is_available():
        return (
            "Poppler is not installed or not on PATH. Install Poppler, or set "
            "POPPLER_PATH to the folder that contains pdftoppm.exe."
        )
    return ""


def tesseract_is_available() -> bool:
    if pytesseract is None:
        return False
    configured = configured_tesseract_path()
    if configured:
        pytesseract.pytesseract.tesseract_cmd = configured
        return True
    return False


def poppler_is_available() -> bool:
    return bool(configured_poppler_path())


def extract_text_with_ocr(path: Path) -> str:
    global LAST_OCR_ERROR
    LAST_OCR_ERROR = ""
    unavailable = ocr_unavailable_message()
    if unavailable:
        LAST_OCR_ERROR = unavailable
        return ""

    convert_options = {
        "dpi": 220,
        "first_page": 1,
        "last_page": 5,
    }
    poppler_path = configured_poppler_path()
    if poppler_path:
        convert_options["poppler_path"] = poppler_path

    try:
        images = convert_from_path(str(path), **convert_options)
    except Exception as exc:
        LAST_OCR_ERROR = (
            "OCR could not convert the scanned PDF pages into images. Install Poppler "
            "and set POPPLER_PATH to its bin folder if it is not on PATH. "
            f"Details: {exc}"
        )
        return ""

    text_parts = []
    try:
        for image in images:
            text_parts.append(pytesseract.image_to_string(image))
    except Exception as exc:
        LAST_OCR_ERROR = (
            "OCR could not read text from the scanned PDF images. Install Tesseract OCR "
            "and set TESSERACT_CMD to tesseract.exe if it is not on PATH. "
            f"Details: {exc}"
        )
        return ""

    return "\n".join(text_parts)


def unreadable_resume_message(extension: str) -> str:
    if extension == "pdf" and LAST_OCR_ERROR:
        return f"Unable to parse readable text from this scanned PDF. {compact_reason(LAST_OCR_ERROR, 260)}"
    if pytesseract is None or convert_from_path is None:
        return (
            "Unable to parse readable text from the resume. If this is a scanned PDF, "
            "install OCR support or export it as DOCX/PDF with selectable text."
        )
    return (
        "Unable to parse readable text from the resume. If this is a scanned PDF, "
        "make sure Tesseract OCR and Poppler are installed correctly, or export it "
        "as DOCX/PDF with selectable text."
    )


def extract_text_from_docx(path: Path) -> str:
    if Document is None:
        raise RuntimeError("DOCX support requires python-docx. Install it with: pip install python-docx")
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def extract_resume_text(path: Path, extension: str) -> str:
    if extension == "pdf":
        text = extract_text_from_pdf(path)
        return text if text.strip() else extract_text_with_ocr(path)
    if extension == "docx":
        return extract_text_from_docx(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip()).lower()


def alias_in_text(text: str, alias: str) -> bool:
    return re.search(rf"(?<![a-z0-9+#.-]){re.escape(alias)}(?![a-z0-9+#.-])", text) is not None


def extract_experience_years(raw_text: str) -> int:
    text = normalize_text(raw_text)
    current_year = utc_now().year
    explicit_patterns = [
        r"\b(\d{1,2})\+?\s*(?:years|yrs)(?:\s+of)?\s+(?:work\s+)?experience\b",
        r"\bexperience(?:\s+of|\s*:)?\s*(\d{1,2})\+?\s*(?:years|yrs)\b",
    ]
    for pattern in explicit_patterns:
        match = re.search(pattern, text)
        if match:
            return min(int(match.group(1)), 25)

    month_names = (
        "jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
        "jul(?:y)?|aug(?:ust)?|sep(?:tember)?|sept|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?"
    )
    date_part = rf"(?:{month_names})?\s*(?:19|20)\d{{2}}"
    range_pattern = re.compile(
        rf"(?P<start>{date_part})\s*(?:-|–|—|to|until)\s*(?P<end>present|current|now|{date_part})"
    )
    work_words = {
        "experience", "employment", "work", "intern", "internship", "company", "developer",
        "engineer", "analyst", "manager", "assistant", "executive", "consultant", "specialist",
    }
    education_words = {"education", "university", "college", "school", "degree", "bachelor", "master", "diploma"}
    intervals = []
    for match in range_pattern.finditer(text):
        start_year = int(re.search(r"(?:19|20)\d{2}", match.group("start")).group(0))
        end_text = match.group("end")
        end_year_match = re.search(r"(?:19|20)\d{2}", end_text)
        end_year = current_year if end_text in {"present", "current", "now"} else int(end_year_match.group(0))
        if end_year < start_year or start_year < 1970 or end_year > current_year:
            continue

        window = text[max(0, match.start() - 80): match.end() + 80]
        has_work_context = any(word in window for word in work_words)
        has_education_context = any(word in window for word in education_words)
        if has_work_context and not (has_education_context and not has_work_context):
            intervals.append((start_year, end_year))

    if not intervals:
        return 0

    covered_years = set()
    for start_year, end_year in intervals:
        covered_years.update(range(start_year, end_year))
    return min(len(covered_years), 25)


def extract_resume_profile(raw_text: str) -> dict:
    text = normalize_text(raw_text)
    doc = NLP(raw_text) if NLP else None
    entities = []
    if doc and hasattr(doc, "ents"):
        entities = [
            {"text": ent.text, "label": ent.label_}
            for ent in doc.ents
            if ent.label_ in {"ORG", "PERSON", "GPE", "DATE", "PRODUCT", "NORP"}
        ][:20]

    skills = []
    for canonical, aliases in SKILL_SYNONYMS.items():
        if any(alias_in_text(text, alias) for alias in aliases):
            skills.append(canonical)

    action_verb_hits = sorted({verb for verb in ACTION_VERBS if re.search(rf"\b{re.escape(verb)}\b", text)})
    weak_verb_hits = sorted({verb for verb in WEAK_VERBS if verb in text})
    education = sorted({
        word for word in ["bachelor", "master", "phd", "degree", "university", "college", "diploma"]
        if word in text
    })

    return {
        "skills": skills,
        "entities": entities,
        "experience_years": extract_experience_years(raw_text),
        "action_verbs": action_verb_hits,
        "weak_verbs": weak_verb_hits,
        "education": education,
        "word_count": len(re.findall(r"\b\w+\b", text)),
    }


def detect_contact_info(text: str) -> int:
    score = 0
    if CONTACT_PATTERN.search(text):
        score += 5
    if "linkedin.com" in text or "portfolio" in text or "github.com" in text:
        score += 5
    return min(score, 10)


def explain_contact_info(text: str, score: int) -> str:
    has_contact = bool(CONTACT_PATTERN.search(text))
    has_profile_link = "linkedin.com" in text or "portfolio" in text or "github.com" in text
    if has_contact and has_profile_link:
        return "Found direct contact details plus a portfolio, LinkedIn, or GitHub signal."
    if has_contact:
        return "Found email or phone details; add LinkedIn, GitHub, or portfolio when relevant."
    if has_profile_link:
        return "Found a profile link, but no clear email or phone number was detected."
    return "No clear email, phone, LinkedIn, GitHub, or portfolio signal was detected."


def detect_summary(text: str) -> int:
    markers = ["summary", "objective", "experience", "achievements", "years of experience", "result-oriented"]
    hits = sum(1 for marker in markers if marker in text)
    return min(10, max(0, hits * 2))


def explain_summary(text: str, score: int) -> str:
    markers = ["summary", "objective", "achievements", "years of experience", "result-oriented"]
    found = [marker for marker in markers if marker in text]
    if score >= 7:
        return f"Detected summary-style signals: {', '.join(found[:4])}."
    if found:
        return f"Found some summary signals ({', '.join(found[:3])}), but the section could be clearer."
    return "No strong professional summary or objective section was detected."


def detect_experience(text: str, profile: dict | None = None) -> int:
    bullets = count_bullet_markers(text)
    role_keywords = len(re.findall(r"\b(?:managed|led|developed|implemented|engineered|designed|built|achieved|improved|created)\b", text))
    experience_years = (profile or {}).get("experience_years", 0)
    if experience_years >= 5 and role_keywords >= 3:
        return 10
    if experience_years >= 2 and role_keywords >= 2:
        return 8
    if bullets >= 4 and role_keywords >= 3:
        return 10
    if bullets >= 2 and role_keywords >= 2:
        return 8
    if bullets >= 1 and role_keywords >= 1:
        return 6
    return 3 if "experience" in text else 0


def explain_experience(text: str, profile: dict | None, score: int) -> str:
    bullets = count_bullet_markers(text)
    role_keywords = len(re.findall(r"\b(?:managed|led|developed|implemented|engineered|designed|built|achieved|improved|created)\b", text))
    experience_years = (profile or {}).get("experience_years", 0)
    if score >= 8:
        return f"Detected {bullets} bullet markers, {role_keywords} impact verbs, and about {experience_years} years of experience."
    return f"Only {bullets} bullet markers and {role_keywords} impact verbs were detected; measurable experience bullets would help."


def detect_education(text: str) -> int:
    degree_words = ["bachelor", "master", "phd", "university", "college", "degree", "graduat"]
    hits = sum(1 for word in degree_words if word in text)
    year_matches = len(YEAR_PATTERN.findall(text))
    score = min(10, hits * 2 + min(year_matches, 2) * 2)
    if "gpa" in text and ("3.5" in text or "3.6" in text or "3.7" in text or "3.8" in text or "4.0" in text):
        score += 2
    return min(score, 10)


def explain_education(text: str, score: int) -> str:
    degree_words = ["bachelor", "master", "phd", "university", "college", "degree", "graduat", "diploma"]
    found = [word for word in degree_words if word in text]
    if found:
        return f"Found education signals: {', '.join(found[:5])}."
    return "No clear degree, school, college, university, or diploma signal was detected."


def detect_skills(text: str, profile: dict | None = None) -> int:
    extracted = set((profile or {}).get("skills", []))
    tech_skills = [
        "python", "sql", "pandas", "excel", "flask", "data visualization",
        "javascript", "machine learning", "nlp", "testing", "api development",
        "git", "html", "css"
    ]
    soft_skills = ["communication", "teamwork", "problem solving", "leadership", "adaptability", "collaboration", "organization"]
    found = len(extracted)
    found += sum(1 for skill in tech_skills + soft_skills if skill in text and skill not in extracted)
    return min(10, found * 2)


def explain_skills(text: str, profile: dict | None, score: int) -> str:
    skills = (profile or {}).get("skills", [])
    if skills:
        return f"Detected explicit skills such as {', '.join(skills[:6])}."
    return "No canonical skill names from the current skill dictionary were detected."


def detect_tas_optimization(text: str) -> int:
    hits = sum(1 for keyword in JOB_KEYWORDS if keyword in text)
    return min(10, hits * 2)


def explain_tas_optimization(text: str, score: int) -> str:
    found = [keyword for keyword in JOB_KEYWORDS if keyword in text]
    if found:
        return f"Matched TAS-friendly terms: {', '.join(found[:6])}."
    return "Few general TAS-friendly keywords were found in the resume text."


def detect_consistency(text: str) -> int:
    date_formats = re.findall(r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4})\b", text)
    bullets = count_bullet_markers(text)
    pattern = 10 if bullets >= 3 and len(date_formats) >= 2 else 6 if bullets >= 2 else 4
    return pattern


def explain_consistency(text: str, score: int) -> str:
    date_formats = re.findall(r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4})\b", text)
    bullets = count_bullet_markers(text)
    return f"Detected {bullets} bullet markers and {len(date_formats)} date markers."


def detect_proofreading(text: str) -> int:
    typo_count = sum(text.count(mistake) for mistake in COMMON_TYPO)
    punctuation_issues = len(re.findall(r"\s{2,}", text))
    score = 10 - min(6, typo_count + punctuation_issues)
    return max(0, score)


def explain_proofreading(text: str, score: int) -> str:
    typos = [mistake for mistake in COMMON_TYPO if mistake in text]
    punctuation_issues = len(re.findall(r"\s{2,}", text))
    if not typos and not punctuation_issues:
        return "No common typo or spacing issues from the local checker were detected."
    issues = []
    if typos:
        issues.append(f"possible typos: {', '.join(typos[:4])}")
    if punctuation_issues:
        issues.append(f"{punctuation_issues} spacing issue(s)")
    return compact_reason("; ".join(issues))


def detect_relevance(text: str, target_text: str = "") -> int:
    if target_text:
        similarity = cosine_similarity(tokenize_for_similarity(text), tokenize_for_similarity(target_text))
        return int(np.clip(round(similarity * 18), 0, 10))
    if "high school" in text and any(word in text for word in ["bachelor", "master", "university", "college"]):
        return 6
    if "high school" in text and "degree" not in text:
        return 4
    return 10


def explain_relevance(text: str, target_text: str, score: int) -> str:
    if target_text:
        similarity = cosine_similarity(tokenize_for_similarity(text), tokenize_for_similarity(target_text))
        return f"Target role similarity is {round(similarity * 100)}% using weighted term overlap."
    return "No target job description was supplied, so relevance used general resume completeness signals."


def build_criteria_scores(text: str, extension: str, profile: dict | None = None, target_text: str = "") -> dict:
    criteria_scores = {
        "Contact Information": detect_contact_info(text),
        "Professional Summary": detect_summary(text),
        "Work Experience": detect_experience(text, profile),
        "Education": detect_education(text),
        "Skills": detect_skills(text, profile),
        "TAS Optimization": detect_tas_optimization(text),
        "Consistency": detect_consistency(text),
        "Proofreading": detect_proofreading(text),
        "File Format": 10 if extension.lower() in {"pdf", "docx"} else 7,
        "Relevance": detect_relevance(text, target_text),
    }
    return criteria_scores


def build_score_details(criteria_scores: dict, text: str, extension: str, profile: dict | None = None, target_text: str = "") -> dict:
    explainers = {
        "Contact Information": explain_contact_info(text, criteria_scores.get("Contact Information", 0)),
        "Professional Summary": explain_summary(text, criteria_scores.get("Professional Summary", 0)),
        "Work Experience": explain_experience(text, profile, criteria_scores.get("Work Experience", 0)),
        "Education": explain_education(text, criteria_scores.get("Education", 0)),
        "Skills": explain_skills(text, profile, criteria_scores.get("Skills", 0)),
        "TAS Optimization": explain_tas_optimization(text, criteria_scores.get("TAS Optimization", 0)),
        "Consistency": explain_consistency(text, criteria_scores.get("Consistency", 0)),
        "Proofreading": explain_proofreading(text, criteria_scores.get("Proofreading", 0)),
        "File Format": f"{extension.upper()} is {'a preferred TAS format' if extension.lower() in {'pdf', 'docx'} else 'accepted but less formatting-stable than PDF or DOCX'}.",
        "Relevance": explain_relevance(text, target_text, criteria_scores.get("Relevance", 0)),
    }
    return {name: compact_reason(reason) for name, reason in explainers.items()}


def calculate_weighted_score(criteria_scores: dict) -> int:
    weighted = sum(criteria_scores[name] * SCORING_WEIGHTS.get(name, 0.0) for name in criteria_scores)
    return int(np.clip(round(weighted * 10), 0, 100))


def scoring_weights_for_role(target_title: str = "") -> dict:
    weights = SCORING_WEIGHTS.copy()
    normalized_title = normalize_text(target_title)
    for role, overrides in ROLE_WEIGHT_OVERRIDES.items():
        if role in normalized_title:
            weights.update(overrides)
            break
    total = sum(weights.values())
    return {name: value / total for name, value in weights.items()} if total else SCORING_WEIGHTS.copy()


def calculate_role_adjusted_score(criteria_scores: dict, target_title: str = "") -> int:
    weights = scoring_weights_for_role(target_title)
    weighted = sum(criteria_scores[name] * weights.get(name, 0.0) for name in criteria_scores)
    return int(np.clip(round(weighted * 10), 0, 100))


def build_analysis_steps(criteria_scores: dict, matched_jobs: list, profile: dict, target_title: str = "") -> list[dict]:
    top_match = matched_jobs[0] if matched_jobs else None
    role_text = target_title or (top_match or {}).get("title") or "the selected role"
    return [
        {
            "label": "Parsed Resume",
            "detail": f"Extracted {profile.get('word_count', 0)} words and {len(profile.get('skills', []))} skill signal(s).",
        },
        {
            "label": "Checked TAS Signals",
            "detail": f"Strongest area: {max(criteria_scores, key=criteria_scores.get)}. Weakest area: {min(criteria_scores, key=criteria_scores.get)}.",
        },
        {
            "label": "Compared Target Role",
            "detail": f"Compared resume terms with {role_text} using role hints and weighted keyword overlap.",
        },
        {
            "label": "Generated Fixes",
            "detail": "Prioritized missing keywords, weak phrasing, measurable impact, and formatting signals.",
        },
    ]


def extract_improvable_bullets(raw_text: str) -> list[str]:
    bullets = []
    for line in raw_text.splitlines():
        cleaned = line.strip().lstrip("-* \u2022").strip()
        if len(cleaned) < 18:
            continue
        lowered = normalize_text(cleaned)
        if any(weak in lowered for weak in WEAK_VERBS) or not re.search(r"\d|%|\$|increased|reduced|improved|optimized|delivered", lowered):
            bullets.append(cleaned)
    return bullets[:3]


def rewrite_resume_bullets(raw_text: str, target_title: str = "") -> list[dict]:
    rewrites = []
    role_phrase = f" for {target_title}" if target_title else ""
    for bullet in extract_improvable_bullets(raw_text):
        cleaned = re.sub(r"\b(?:helped|assisted|worked on|responsible for|handled|participated in)\b", "Delivered", bullet, flags=re.IGNORECASE)
        cleaned = cleaned[0].upper() + cleaned[1:] if cleaned else bullet
        if not re.search(r"\d|%|\$|increased|reduced|improved|optimized", normalize_text(cleaned)):
            cleaned = f"Improved {cleaned[0].lower() + cleaned[1:]}{role_phrase}, adding measurable outcomes such as time saved, quality gains, or user impact."
        rewrites.append({"before": bullet, "after": compact_reason(cleaned, 240)})
    return rewrites


def text_contains_alias(text: str, alias: str) -> bool:
    return alias_in_text(text, alias)


def tokenize_for_similarity(text: str) -> list[str]:
    tokens = re.findall(r"\b[a-zA-Z][a-zA-Z+#.-]{1,}\b", normalize_text(text))
    stop_words = {
        "and", "the", "for", "with", "from", "that", "this", "into", "your", "you",
        "are", "will", "have", "has", "our", "resume", "experience", "work"
    }
    expanded = [token for token in tokens if token not in stop_words]
    normalized = normalize_text(text)
    for canonical, aliases in SKILL_SYNONYMS.items():
        if any(text_contains_alias(normalized, alias) for alias in aliases):
            expanded.extend(canonical.split())
    return expanded


def build_weighted_terms(text: str) -> Counter:
    tokens = tokenize_for_similarity(text)
    counts = Counter(tokens)
    normalized = normalize_text(text)
    for canonical, aliases in SKILL_SYNONYMS.items():
        if any(text_contains_alias(normalized, alias) for alias in aliases):
            counts[canonical] += 3
    for phrase in re.findall(r"\b(?:required|must have|proficient in|experience with|responsible for)\s+([^.;:\n]+)", normalized):
        for token in tokenize_for_similarity(phrase):
            counts[token] += 2
    return counts


def infer_role_keywords(title: str) -> list[str]:
    normalized_title = normalize_text(title)
    inferred: list[str] = []
    for role, keywords in ROLE_KEYWORD_HINTS.items():
        if role in normalized_title:
            inferred.extend(keywords)
    return list(dict.fromkeys(inferred))


def keyword_found_in_text(keyword: str, normalized_text: str) -> bool:
    aliases = SKILL_SYNONYMS.get(keyword, [keyword])
    return any(text_contains_alias(normalized_text, alias) for alias in aliases)


def weighted_cosine_similarity(left: Counter, right: Counter) -> float:
    if not left or not right:
        return 0.0
    shared = set(left) & set(right)
    numerator = sum(left[token] * right[token] for token in shared)
    left_norm = math.sqrt(sum(value * value for value in left.values()))
    right_norm = math.sqrt(sum(value * value for value in right.values()))
    return numerator / (left_norm * right_norm) if left_norm and right_norm else 0.0


def cosine_similarity(left_tokens: list[str], right_tokens: list[str]) -> float:
    left = Counter(left_tokens)
    right = Counter(right_tokens)
    if not left or not right:
        return 0.0
    shared = set(left) & set(right)
    numerator = sum(left[token] * right[token] for token in shared)
    left_norm = math.sqrt(sum(value * value for value in left.values()))
    right_norm = math.sqrt(sum(value * value for value in right.values()))
    return numerator / (left_norm * right_norm) if left_norm and right_norm else 0.0


def build_target_job_match(text: str, title: str, description: str) -> dict | None:
    target_text = f"{title} {description}".strip()
    if not target_text:
        return None
    inferred_keywords = infer_role_keywords(title)
    expanded_target_text = f"{target_text} {' '.join(inferred_keywords)}".strip()
    resume_terms = build_weighted_terms(text)
    target_counter = build_weighted_terms(expanded_target_text)
    candidate_keywords = [
        token for token, _ in target_counter.most_common()
        if len(token) > 2 and token not in LOW_SIGNAL_ROLE_WORDS
    ][:18]
    normalized_resume = normalize_text(text)
    matched_keywords = [keyword for keyword in candidate_keywords if keyword_found_in_text(keyword, normalized_resume)]
    missing_keywords = [
        keyword for keyword in candidate_keywords
        if not keyword_found_in_text(keyword, normalized_resume)
    ][:8]
    similarity = weighted_cosine_similarity(resume_terms, target_counter)
    coverage = len(matched_keywords) / max(len(candidate_keywords), 1)
    match_score = (similarity * 0.6) + (coverage * 0.4)
    return {
        "title": title or "Target Role",
        "description": description,
        "match_score": int(np.clip(round(match_score * 100), 0, 100)),
        "matched_keywords": matched_keywords[:10],
        "missing_keywords": missing_keywords,
        "source": "target",
        "explanation": f"Matched {len(matched_keywords)} of {len(candidate_keywords)} high-signal target terms.",
    }


def find_applicable_jobs(text: str, profile: dict | None = None, target_job: dict | None = None) -> list:
    jobs_df = load_job_data()
    text_lower = text.lower()
    resume_terms = build_weighted_terms(text)
    recommendations = []
    if target_job:
        recommendations.append(target_job)
    for _, row in jobs_df.iterrows():
        keywords = [kw.strip() for kw in str(row["keywords"]).split(",") if kw.strip()]
        matches = sum(1 for kw in keywords if kw in text_lower)
        keyword_ratio = matches / max(len(keywords), 1)
        job_text = f"{row['title']} {row.get('description', '')} {' '.join(keywords)}"
        similarity = weighted_cosine_similarity(resume_terms, build_weighted_terms(job_text))
        combined_score = (keyword_ratio * 0.55) + (similarity * 0.45)
        matched_keywords = [kw for kw in keywords if kw in text_lower]
        missing_keywords = [kw for kw in keywords if kw not in text_lower][:5]
        if combined_score >= 0.18 or matched_keywords:
            recommendations.append({
                "title": row["title"],
                "description": row.get("description", ""),
                "match_score": int(np.clip(round(combined_score * 100), 0, 100)),
                "matched_keywords": matched_keywords,
                "missing_keywords": missing_keywords,
                "source": "catalog",
            })
    recommendations.sort(key=lambda item: (item.get("source") == "target", item["match_score"]), reverse=True)
    return recommendations[:5]


def build_resume_suggestions(criteria_scores: dict, matched_jobs: list, profile: dict, text: str) -> list[str]:
    suggestions = []
    skills = set(profile.get("skills", []))
    missing_from_top_jobs = []
    for job in matched_jobs[:3]:
        missing_from_top_jobs.extend(job.get("missing_keywords", []))
    most_common_missing = [skill for skill, _ in Counter(missing_from_top_jobs).most_common(5)]

    if criteria_scores.get("Contact Information", 0) < 8:
        suggestions.append("Add a clear email/phone line and include LinkedIn or GitHub when relevant.")
    if criteria_scores.get("Professional Summary", 0) < 7:
        suggestions.append("Add a 2-3 line professional summary tailored to the target role.")
    if criteria_scores.get("Work Experience", 0) < 7:
        suggestions.append("Rewrite experience bullets with impact: action verb + task + measurable outcome.")
    if criteria_scores.get("Skills", 0) < 8 and most_common_missing:
        suggestions.append(f"Consider adding relevant skills from matching jobs: {', '.join(most_common_missing)}.")
    elif criteria_scores.get("Skills", 0) < 8:
        suggestions.append("Add a dedicated skills section with tools, methods, and soft skills.")
    if profile.get("weak_verbs"):
        suggestions.append(f"Replace weak phrasing like {', '.join(profile['weak_verbs'][:3])} with stronger action verbs.")
    if len(profile.get("action_verbs", [])) < 4:
        suggestions.append("Use more accomplishment verbs such as built, improved, automated, delivered, or optimized.")
    if profile.get("word_count", 0) < 180:
        suggestions.append("The resume text is quite short; add project, role, education, and achievement detail.")
    if not skills:
        suggestions.append("Add explicit skill names so TAS systems can match your profile reliably.")
    target_match = next((job for job in matched_jobs if job.get("source") == "target"), None)
    if target_match and target_match.get("missing_keywords"):
        suggestions.append(f"Tailor this resume to the target role by adding evidence for: {', '.join(target_match['missing_keywords'][:5])}.")
    return suggestions[:6]


def save_report(
    user_id: int | None,
    filename: str,
    score: int,
    matched_jobs: list,
    criteria_scores: dict,
    suggestions: list[str] | None = None,
    analysis: dict | None = None,
) -> int:
    conn = sqlite3.connect(DB_PATH)
    try:
        cursor = conn.execute(
            "INSERT INTO reports (user_id, filename, uploaded_at, score, matched_jobs, criteria, suggestions, analysis) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (
                user_id,
                filename,
                utc_timestamp(),
                score,
                json.dumps(matched_jobs),
                json.dumps(criteria_scores),
                json.dumps(suggestions or []),
                json.dumps(analysis or {}),
            ),
        )
        conn.commit()
        return int(cursor.lastrowid)
    finally:
        conn.close()


def save_contact_submission(user_id: int | None, name: str, email: str, subject: str, message: str) -> None:
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(
            "INSERT INTO contacts (user_id, name, email, subject, message, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (user_id, name, email, subject, message, utc_timestamp()),
        )
        conn.commit()
    finally:
        conn.close()


def get_latest_score(user_id: int | None) -> int | None:
    conn = sqlite3.connect(DB_PATH)
    try:
        if user_id is None:
            row = conn.execute(
                "SELECT score FROM reports WHERE user_id IS NULL ORDER BY uploaded_at DESC LIMIT 1"
            ).fetchone()
        else:
            row = conn.execute(
                "SELECT score FROM reports WHERE user_id = ? ORDER BY uploaded_at DESC LIMIT 1",
                (user_id,),
            ).fetchone()
        return row[0] if row else None
    finally:
        conn.close()


def get_user(username: str) -> dict | None:
    conn = sqlite3.connect(DB_PATH)
    try:
        row = conn.execute("SELECT id, username, password_hash FROM users WHERE username = ?", (username,)).fetchone()
        if row:
            return {"id": row[0], "username": row[1], "password_hash": row[2]}
        return None
    finally:
        conn.close()


def get_current_user() -> dict | None:
    username = session.get("user")
    if not username:
        return None
    return get_user(username)


def create_user(username: str, password: str) -> bool:
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(
            "INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
            (username, generate_password_hash(password), utc_timestamp()),
        )
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()


def ensure_user_columns(conn: sqlite3.Connection) -> None:
    columns = [row[1] for row in conn.execute("PRAGMA table_info(users)").fetchall()]
    if "created_at" not in columns:
        conn.execute("ALTER TABLE users ADD COLUMN created_at TEXT")
    conn.commit()


def ensure_contact_columns(conn: sqlite3.Connection) -> None:
    columns = [row[1] for row in conn.execute("PRAGMA table_info(contacts)").fetchall()]
    if "user_id" not in columns:
        conn.execute("ALTER TABLE contacts ADD COLUMN user_id INTEGER")
    conn.commit()


def ensure_report_columns(conn: sqlite3.Connection) -> None:
    columns = [row[1] for row in conn.execute("PRAGMA table_info(reports)").fetchall()]
    for column, column_type in {"suggestions": "TEXT", "analysis": "TEXT"}.items():
        if column not in columns:
            conn.execute(f"ALTER TABLE reports ADD COLUMN {column} {column_type}")
    conn.commit()


def ensure_reports_user_scope(conn: sqlite3.Connection) -> None:
    columns = [row[1] for row in conn.execute("PRAGMA table_info(reports)").fetchall()]
    if "user_id" not in columns:
        conn.execute("ALTER TABLE reports ADD COLUMN user_id INTEGER")
        conn.commit()


def assign_legacy_reports_to_default_user(conn: sqlite3.Connection) -> None:
    primary_row = get_primary_history_user_row(conn)
    if primary_row:
        conn.execute(
            "UPDATE reports SET user_id = ? WHERE user_id IS NULL",
            (primary_row[0],),
        )
        conn.commit()


def ensure_runtime_initialized() -> None:
    global _DB_READY
    if _DB_READY:
        return
    init_db()
    load_job_data()
    _DB_READY = True


@app.before_request
def initialize_runtime_on_request() -> None:
    ensure_runtime_initialized()


def normalize_report_payload(row) -> dict | None:
    if not row:
        return None
    matched_jobs = parse_report_payload(row[4], [])
    criteria_scores = parse_report_payload(row[5], {})
    suggestions = parse_report_payload(row[6], [])
    analysis = parse_report_payload(row[7], {})
    if isinstance(matched_jobs, str):
        matched_jobs = [{"title": title.strip()} for title in matched_jobs.split(",") if title.strip()]
    return {
        "report": {
            "id": row[0],
            "filename": row[1],
            "uploaded_at": row[2],
            "score": row[3],
        },
        "matched_jobs": matched_jobs if isinstance(matched_jobs, list) else [],
        "criteria_scores": criteria_scores if isinstance(criteria_scores, dict) else {},
        "suggestions": suggestions if isinstance(suggestions, list) else [],
        "analysis": analysis if isinstance(analysis, dict) else {},
    }


def get_report_payload(report_id: int, user_id: int) -> dict | None:
    conn = sqlite3.connect(DB_PATH)
    try:
        accessible_user_ids = get_accessible_user_ids(conn, user_id)
        placeholders = ", ".join(["?"] * len(accessible_user_ids))
        row = conn.execute(
            f"SELECT id, filename, uploaded_at, score, matched_jobs, criteria, suggestions, analysis "
            f"FROM reports WHERE id = ? AND user_id IN ({placeholders})",
            (report_id, *accessible_user_ids),
        ).fetchone()
    finally:
        conn.close()
    return normalize_report_payload(row)


def escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_simple_pdf(lines: list[str]) -> bytes:
    content = ["BT", "/F1 12 Tf", "50 780 Td", "16 TL"]
    for line in lines[:42]:
        content.append(f"({escape_pdf_text(line[:95])}) Tj")
        content.append("T*")
    content.append("ET")
    stream = "\n".join(content).encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    output = BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(output.tell())
        output.write(f"{index} 0 obj\n".encode("ascii"))
        output.write(obj)
        output.write(b"\nendobj\n")
    xref = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.write(f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode("ascii"))
    return output.getvalue()


def report_lines(payload: dict) -> list[str]:
    report = payload["report"]
    analysis = payload.get("analysis", {})
    score_details = analysis.get("score_details", {})
    lines = [
        "Talent Acquisition System Report",
        f"File: {report['filename']}",
        f"Saved: {report['uploaded_at']}",
        f"Score: {report['score']} / 100",
        "",
        "Criteria",
    ]
    for name, value in payload["criteria_scores"].items():
        lines.append(f"- {name}: {value}/10")
        if score_details.get(name):
            lines.append(f"  {score_details[name]}")
    lines.append("")
    lines.append("Suggestions")
    lines.extend(f"- {suggestion}" for suggestion in payload["suggestions"] or ["No suggestions saved."])
    if analysis.get("bullet_rewrites"):
        lines.append("")
        lines.append("Suggested Bullet Rewrites")
        for rewrite in analysis["bullet_rewrites"][:3]:
            lines.append(f"- Before: {rewrite.get('before', '')}")
            lines.append(f"  After: {rewrite.get('after', '')}")
    if analysis.get("analysis_steps"):
        lines.append("")
        lines.append("Analysis Steps")
        for step in analysis["analysis_steps"][:4]:
            lines.append(f"- {step.get('label', '')}: {step.get('detail', '')}")
    lines.append("")
    lines.append("Matches")
    for job in payload["matched_jobs"]:
        if isinstance(job, dict):
            lines.append(f"- {job.get('title', 'Role')}: {job.get('match_score', 0)}%")
    return lines


def login_required(view_func):
    @wraps(view_func)
    def wrapped(*args, **kwargs):
        if session.get("user") is None:
            flash("Please log in to access that page.", "error")
            return redirect(url_for("login"))
        return view_func(*args, **kwargs)
    return wrapped


def admin_required(view_func):
    @wraps(view_func)
    def wrapped(*args, **kwargs):
        current_user = get_current_user()
        if not current_user:
            flash("Please log in to access that page.", "error")
            return redirect(url_for("login"))
        if current_user["username"] != app.config.get("ADMIN_USERNAME", "admin"):
            flash("Admin access is required.", "error")
            return redirect(url_for("index"))
        return view_func(*args, **kwargs)
    return wrapped


@app.route("/login", methods=["GET", "POST"])
@rate_limit(10, 300)
def login():
    if session.get("user"):
        flash("You are already logged in.", "success")
        return redirect(url_for("index"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        user = get_user(username)
        if user and check_password_hash(user["password_hash"], password):
            session["user"] = user["username"]
            flash(f"Welcome back, {user['username']}.", "success")
            return redirect(url_for("index"))
        flash("Invalid username or password.", "error")
        return redirect(url_for("login"))
    return render_template("login.html")


@app.route("/register", methods=["GET", "POST"])
@rate_limit(5, 300)
def register():
    if session.get("user"):
        flash("You are already logged in.", "success")
        return redirect(url_for("index"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        confirm_password = request.form.get("confirm_password", "")

        if len(username) < 3:
            flash("Username must be at least 3 characters long.", "error")
            return redirect(url_for("register"))
        if not re.fullmatch(r"[A-Za-z0-9_.-]+", username):
            flash("Username can only contain letters, numbers, dots, underscores, and hyphens.", "error")
            return redirect(url_for("register"))
        if len(password) < 8:
            flash("Password must be at least 8 characters long.", "error")
            return redirect(url_for("register"))
        if password != confirm_password:
            flash("Passwords do not match.", "error")
            return redirect(url_for("register"))
        if get_user(username):
            flash("That username is already registered. Please choose another one.", "error")
            return redirect(url_for("register"))

        if create_user(username, password):
            flash("Account created successfully. Please log in with your new credentials.", "success")
            return redirect(url_for("login"))

        flash("We could not create your account right now. Please try again.", "error")
        return redirect(url_for("register"))

    return render_template("register.html")


@app.route("/logout")
def logout():
    session.pop("user", None)
    flash("Logged out successfully.", "success")
    return redirect(url_for("index"))


@app.route("/contact", methods=["GET", "POST"])
@rate_limit(5, 300)
def contact():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip()
        subject = request.form.get("subject", "").strip()
        message = request.form.get("message", "").strip()
        if not name or not email or not message:
            flash("Please fill in your name, email, and message.", "error")
            return redirect(url_for("contact"))
        if not EMAIL_PATTERN.fullmatch(email):
            flash("Please enter a valid email address.", "error")
            return redirect(url_for("contact"))
        if any(len(value) > MAX_TEXT_INPUT_LENGTH for value in [name, email, subject, message]):
            flash("Please keep contact fields under 5,000 characters.", "error")
            return redirect(url_for("contact"))
        current_user = get_current_user()
        save_contact_submission(current_user["id"] if current_user else None, name, email, subject, message)
        flash("Your message has been sent. You will be contacted shortly.", "success")
        return redirect(url_for("contact"))
    return render_template("contact.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/", methods=["GET", "POST"])
@rate_limit(12, 300)
def index():
    if request.method == "POST":
        resume_file = request.files.get("resume_file")
        if not resume_file or resume_file.filename == "":
            flash("Please upload a resume file in PDF, DOCX, or TXT format.", "error")
            return redirect(url_for("index"))

        filename = secure_filename(resume_file.filename)
        if not allowed_file(filename):
            flash("Only PDF, DOCX, and TXT resume uploads are supported.", "error")
            return redirect(url_for("index"))

        current_user = get_current_user()
        save_path = None
        try:
            file_ext = filename.rsplit(".", 1)[1].lower()
            target_title = request.form.get("job_title", "").strip()
            target_description = request.form.get("job_description", "").strip()
            if len(target_title) > 160 or len(target_description) > MAX_TEXT_INPUT_LENGTH:
                flash("Please keep the target role title and description shorter before scanning.", "error")
                return redirect(url_for("index"))
            target_text = f"{target_title} {target_description}".strip()
            save_path = save_upload_temporarily(resume_file, filename, current_user["id"] if current_user else None)
            if not file_signature_matches(save_path, file_ext):
                flash("The uploaded file content does not match its extension. Please upload a valid PDF, DOCX, or TXT resume.", "error")
                return redirect(url_for("index"))
            extracted_text = extract_resume_text(save_path, file_ext)

            if not extracted_text.strip():
                flash(unreadable_resume_message(file_ext), "error")
                return redirect(url_for("index"))

            text = normalize_text(extracted_text)
            profile = extract_resume_profile(extracted_text)
            target_job = build_target_job_match(text, target_title, target_description)
            criteria_scores = build_criteria_scores(text, file_ext, profile, normalize_text(target_text))
            score_details = build_score_details(criteria_scores, text, file_ext, profile, normalize_text(target_text))
            total_score = calculate_role_adjusted_score(criteria_scores, target_title)
            matched_jobs = find_applicable_jobs(text, profile, target_job)
            suggestions = build_resume_suggestions(criteria_scores, matched_jobs, profile, text)
            previous_score = get_latest_score(current_user["id"] if current_user else None)
            analysis_steps = build_analysis_steps(criteria_scores, matched_jobs, profile, target_title)
            bullet_rewrites = rewrite_resume_bullets(extracted_text, target_title)
            profile["target_role"] = target_title
            profile["target_description_provided"] = bool(target_description)
            profile["score_details"] = score_details
            profile["score_weights"] = scoring_weights_for_role(target_title)
            profile["analysis_steps"] = analysis_steps
            profile["bullet_rewrites"] = bullet_rewrites
            report_id = save_report(
                current_user["id"] if current_user else None,
                filename,
                total_score,
                matched_jobs,
                criteria_scores,
                suggestions,
                profile,
            )

            return render_template(
                "result.html",
                filename=filename,
                score=total_score,
                criteria_scores=criteria_scores,
                score_details=score_details,
                analysis_steps=analysis_steps,
                bullet_rewrites=bullet_rewrites,
                matched_jobs=matched_jobs,
                suggestions=suggestions,
                analysis=profile,
                previous_score=previous_score,
                target_title=target_title,
                report_id=report_id,
            )
        except Exception as e:
            traceback.print_exc()
            flash(f"An unexpected error occurred while processing the resume: {e}", "error")
            return redirect(url_for("index"))
        finally:
            if save_path and save_path.exists():
                save_path.unlink(missing_ok=True)

    return render_template("index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    return index()


@app.route("/history")
@login_required
def history():
    current_user = get_current_user()
    conn = sqlite3.connect(DB_PATH)
    try:
        accessible_user_ids = get_accessible_user_ids(conn, current_user["id"])
        placeholders = ", ".join(["?"] * len(accessible_user_ids))
        raw_rows = conn.execute(
            f"SELECT id, filename, uploaded_at, score, matched_jobs "
            f"FROM reports WHERE user_id IN ({placeholders}) ORDER BY uploaded_at DESC LIMIT 20",
            tuple(accessible_user_ids),
        ).fetchall()
    finally:
        conn.close()
    rows = [
        {
            "id": row[0],
            "filename": row[1],
            "uploaded_at": row[2],
            "score": row[3],
            "matched_jobs_summary": summarize_matched_jobs(row[4]),
        }
        for row in raw_rows
    ]
    scores = [row["score"] for row in rows if row["score"] is not None]
    dashboard = {
        "total_scans": len(rows),
        "average_score": round(sum(scores) / len(scores)) if scores else 0,
        "best_score": max(scores) if scores else 0,
        "trend_scores": list(reversed(scores[:8])),
        "trend_labels": [f"Scan {index + 1}" for index in range(min(len(scores), 8))],
    }
    return render_template("history.html", rows=rows, dashboard=dashboard)


@app.route("/history/<int:report_id>")
@login_required
def history_detail(report_id: int):
    current_user = get_current_user()
    payload = get_report_payload(report_id, current_user["id"])

    if not payload:
        flash("That history entry is not available for this account.", "error")
        return redirect(url_for("history"))

    return render_template("history_detail.html", **payload)


@app.route("/history/<int:report_id>/export.pdf")
@login_required
def export_report_pdf(report_id: int):
    current_user = get_current_user()
    payload = get_report_payload(report_id, current_user["id"])
    if not payload:
        flash("That history entry is not available for this account.", "error")
        return redirect(url_for("history"))
    pdf = build_simple_pdf(report_lines(payload))
    filename = secure_filename(f"tas-report-{report_id}.pdf")
    return Response(
        pdf,
        mimetype="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.route("/admin", methods=["GET", "POST"])
@admin_required
@rate_limit(10, 300)
def admin():
    if request.method == "POST":
        title = request.form.get("title", "").strip()
        keywords = request.form.get("keywords", "").strip()
        description = request.form.get("description", "").strip()
        if not title or not keywords:
            flash("Job title and keywords are required.", "error")
            return redirect(url_for("admin"))
        if any(len(value) > MAX_CATALOG_FIELD_LENGTH for value in [title, keywords, description]):
            flash("Please keep catalog fields under 800 characters.", "error")
            return redirect(url_for("admin"))
        jobs_df = load_job_data()
        new_row = pd.DataFrame([{"title": title, "keywords": keywords, "description": description}])
        jobs_df = pd.concat([jobs_df, new_row], ignore_index=True)
        jobs_df.to_csv(JOB_DATA_PATH, index=False)
        flash("Job role added to the recommendation catalog.", "success")
        return redirect(url_for("admin"))

    conn = sqlite3.connect(DB_PATH)
    try:
        contacts = conn.execute(
            "SELECT contacts.name, contacts.email, contacts.subject, contacts.message, contacts.created_at, users.username "
            "FROM contacts LEFT JOIN users ON contacts.user_id = users.id "
            "ORDER BY contacts.created_at DESC LIMIT 20"
        ).fetchall()
        report_count = conn.execute("SELECT COUNT(*) FROM reports").fetchone()[0]
        user_count = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    finally:
        conn.close()
    jobs = load_job_data().to_dict(orient="records")
    return render_template(
        "admin.html",
        contacts=contacts,
        jobs=jobs,
        report_count=report_count,
        user_count=user_count,
    )


if __name__ == "__main__":
    ensure_runtime_initialized()
    app.run(
        debug=os.environ.get("FLASK_DEBUG") == "1",
        host=os.environ.get("FLASK_RUN_HOST", "127.0.0.1"),
        port=int(os.environ.get("PORT", "5000")),
    )
